"""Word document report generation."""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from src.calculators.metrics import FinancialMetrics
from src.calculators.ratios import FinancialRatios
from src.calculators.verification import VerificationResult
from src.fetchers.yahoo import CompanyInfo, CorporateAction
from src.models.extraction import (
    ExtractionSession,
    INCOME_STATEMENT_ITEMS,
    BALANCE_SHEET_ITEMS,
    CASH_FLOW_ITEMS,
)


# Colors
HEADER_BLUE = "4472C4"
GREEN_BG = "C6EFCE"
RED_BG = "FFC7CE"
LIGHT_GRAY = "F2F2F2"


def add_formatted_paragraph(doc: Document, text: str, level: int = None):
    """
    Add a paragraph with markdown formatting support.
    Supports:
    - ## Heading (level 2)
    - ### Heading (level 3)
    - **bold**
    - *italic*
    """
    # Check if this is a heading
    heading_match = re.match(r'^(#{1,6})\s+(.+)$', text)
    if heading_match:
        heading_level = len(heading_match.group(1))
        heading_text = heading_match.group(2)
        doc.add_heading(heading_text, level=heading_level)
        return

    # Regular paragraph with inline formatting
    para = doc.add_paragraph()

    # Parse inline formatting: **bold** and *italic*
    # Process **bold** first to avoid conflicts with *italic*
    pattern = r'(\*\*[^*]+\*\*)|(\*[^*]+\*)|([^*]+)'

    for match in re.finditer(pattern, text):
        content = match.group(0)
        if content.startswith('**') and content.endswith('**'):
            # Bold text
            run = para.add_run(content[2:-2])
            run.bold = True
        elif content.startswith('*') and content.endswith('*'):
            # Italic text
            run = para.add_run(content[1:-1])
            run.italic = True
        else:
            # Regular text
            para.add_run(content)


def set_cell_shading(cell, color_hex: str):
    """Set background color of a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def format_currency(value: float, unit: str = "dollars") -> str:
    """
    Format a currency value based on the unit.

    Args:
        value: The numeric value
        unit: The unit of the value ("dollars", "thousands", "millions", etc.)
    """
    if value == 0:
        return "-"

    # Convert value to base dollars based on unit
    unit_lower = unit.lower()
    if "million" in unit_lower:
        value_in_dollars = value * 1e6
    elif "thousand" in unit_lower:
        value_in_dollars = value * 1e3
    else:  # dollars or unknown
        value_in_dollars = value

    # Format the result
    abs_val = abs(value_in_dollars)
    if abs_val >= 1e9:
        formatted = f"${abs_val / 1e9:,.1f} B"
    elif abs_val >= 1e6:
        formatted = f"${abs_val / 1e6:,.1f} M"
    else:
        formatted = f"${abs_val:,.0f}"

    return f"-{formatted}" if value_in_dollars < 0 else formatted


def format_percentage(value: float) -> str:
    """Format a decimal as percentage."""
    if value == 0:
        return "-"
    return f"{value * 100:.2f}%"


def format_ratio(value: float) -> str:
    """Format a ratio value."""
    if value == 0:
        return "-"
    return f"{value:.2f}x"


def format_delta(current: float, prior: float, unit: str = "dollars", is_percentage: bool = False, is_ratio: bool = False) -> tuple[str, str]:
    """
    Format a delta value and return (formatted_string, color).
    Color is 'green', 'red', or 'none'.
    """
    if current == 0 and prior == 0:
        return "-", "none"

    delta = current - prior

    if delta == 0:
        return "-", "none"

    # Determine color based on whether increase is good
    # For most metrics, increase is good (green)
    # For some ratios like Net Debt/EBITDA, decrease is good
    color = "green" if delta > 0 else "red"

    if is_percentage:
        formatted = f"{delta * 100:+.2f}%"
    elif is_ratio:
        formatted = f"{delta:+.2f}x"
    else:
        # Currency - convert to base dollars first
        unit_lower = unit.lower()
        if "million" in unit_lower:
            delta_in_dollars = delta * 1e6
        elif "thousand" in unit_lower:
            delta_in_dollars = delta * 1e3
        else:
            delta_in_dollars = delta

        abs_delta = abs(delta_in_dollars)
        if abs_delta >= 1e9:
            formatted = f"${abs_delta / 1e9:,.1f} B"
        else:
            formatted = f"${abs_delta / 1e6:,.1f} M"
        if delta_in_dollars < 0:
            formatted = f"-{formatted}"

    return formatted, color


def add_financial_overview_table(doc: Document, metrics: FinancialMetrics, fiscal_year: str, prior_year: str, unit: str = "dollars"):
    """Add the Financial Statements Overview table."""
    doc.add_heading("Financial Statements Overview", level=2)

    # Define rows: (label, current_value, prior_value, is_percentage, is_currency)
    rows_data = [
        ("Tangible Net Worth", metrics.tangible_net_worth, metrics.tangible_net_worth_prior, False, True),
        ("Cash Balance", metrics.cash_balance, metrics.cash_balance_prior, False, True),
        ("Top Line Revenue", metrics.top_line_revenue, metrics.top_line_revenue_prior, False, True),
        ("Gross Profit", metrics.gross_profit, metrics.gross_profit_prior, False, True),
        ("Gross Profit Margin", metrics.gross_profit_margin, metrics.gross_profit_margin_prior, True, False),
        ("Operating Income", metrics.operating_income, metrics.operating_income_prior, False, True),
        ("Operating Income Margin", metrics.operating_income_margin, metrics.operating_income_margin_prior, True, False),
        ("EBITDA", metrics.ebitda, metrics.ebitda_prior, False, True),
        ("EBITDA Margin", metrics.ebitda_margin, metrics.ebitda_margin_prior, True, False),
        ("Adjusted EBITDA", metrics.adjusted_ebitda, metrics.adjusted_ebitda_prior, False, True),
        ("Adj. EBITDA Margin", metrics.adjusted_ebitda_margin, metrics.adjusted_ebitda_margin_prior, True, False),
        ("Net Income", metrics.net_income, metrics.net_income_prior, False, True),
        ("Net Income Margin", metrics.net_income_margin, metrics.net_income_margin_prior, True, False),
    ]

    # Create table
    table = doc.add_table(rows=1 + len(rows_data), cols=4)
    table.style = "Table Grid"

    # Header row
    header_cells = table.rows[0].cells
    headers = ["Financial Statements Overview", fiscal_year[:4], prior_year[:4], "Delta"]
    for i, header in enumerate(headers):
        header_cells[i].text = header
        set_cell_shading(header_cells[i], HEADER_BLUE)
        run = header_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for row_idx, (label, current, prior, is_pct, is_currency) in enumerate(rows_data):
        row = table.rows[row_idx + 1]
        cells = row.cells

        # Label
        cells[0].text = label

        # Current value
        if is_pct:
            cells[1].text = format_percentage(current)
        elif is_currency:
            cells[1].text = format_currency(current, unit=unit)
        else:
            cells[1].text = str(current)

        # Prior value
        if is_pct:
            cells[2].text = format_percentage(prior)
        elif is_currency:
            cells[2].text = format_currency(prior, unit=unit)
        else:
            cells[2].text = str(prior)

        # Delta
        delta_str, color = format_delta(current, prior, unit=unit, is_percentage=is_pct)
        cells[3].text = delta_str

        # Color the delta cell
        if color == "green":
            set_cell_shading(cells[3], GREEN_BG)
        elif color == "red":
            set_cell_shading(cells[3], RED_BG)

        # Alternate row shading
        if row_idx % 2 == 1:
            for cell in [cells[0], cells[1], cells[2]]:
                set_cell_shading(cell, LIGHT_GRAY)

    doc.add_paragraph()  # Spacing


def add_ratios_table(doc: Document, ratios: FinancialRatios, fiscal_year: str, prior_year: str, unit: str = "dollars"):
    """Add the Ratios table."""
    doc.add_heading("Ratios", level=2)

    # For ratios, some decreases are good (debt ratios), some increases are good (coverage)
    # We'll mark which ones "decrease is good"
    rows_data = [
        ("Current Ratio", ratios.current_ratio, ratios.current_ratio_prior, False, False),
        ("Cash Ratio", ratios.cash_ratio, ratios.cash_ratio_prior, False, False),
        ("Debt-to-Equity Ratio", ratios.debt_to_equity, ratios.debt_to_equity_prior, False, True),  # decrease is good
        ("EBITDA Interest Coverage", ratios.ebitda_interest_coverage, ratios.ebitda_interest_coverage_prior, False, False),
        ("Net Debt / EBITDA", ratios.net_debt_to_ebitda, ratios.net_debt_to_ebitda_prior, False, True),  # decrease is good
        ("Net Debt / Adj. EBITDA", ratios.net_debt_to_adj_ebitda, ratios.net_debt_to_adj_ebitda_prior, False, True),  # decrease is good
        ("Days Sales Outstanding", ratios.days_sales_outstanding, ratios.days_sales_outstanding_prior, False, True),  # decrease is good
        ("Working Capital", ratios.working_capital, ratios.working_capital_prior, True, False),  # currency
        ("Return on Assets", ratios.return_on_assets, ratios.return_on_assets_prior, False, False),
        ("Return on Equity", ratios.return_on_equity, ratios.return_on_equity_prior, False, False),
    ]

    table = doc.add_table(rows=1 + len(rows_data), cols=4)
    table.style = "Table Grid"

    # Header row
    header_cells = table.rows[0].cells
    headers = ["Ratios", fiscal_year[:4], prior_year[:4], "Delta"]
    for i, header in enumerate(headers):
        header_cells[i].text = header
        set_cell_shading(header_cells[i], HEADER_BLUE)
        run = header_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for row_idx, (label, current, prior, is_currency, decrease_is_good) in enumerate(rows_data):
        row = table.rows[row_idx + 1]
        cells = row.cells

        # Label
        cells[0].text = label

        # Current/Prior values
        if is_currency:
            cells[1].text = format_currency(current, unit=unit)
            cells[2].text = format_currency(prior, unit=unit)
        elif "Days" in label:
            cells[1].text = f"{current:.1f}" if current else "-"
            cells[2].text = f"{prior:.1f}" if prior else "-"
        else:
            cells[1].text = format_ratio(current)
            cells[2].text = format_ratio(prior)

        # Delta
        delta = current - prior
        if is_currency:
            delta_str, color = format_delta(current, prior, unit=unit)
        elif "Days" in label:
            delta_str = f"{delta:+.1f}" if delta != 0 else "-"
            color = "red" if delta > 0 else "green" if delta < 0 else "none"
        else:
            delta_str = f"{delta:+.2f}x" if delta != 0 else "-"
            color = "green" if delta > 0 else "red" if delta < 0 else "none"

        # Flip color for metrics where decrease is good
        if decrease_is_good and color != "none":
            color = "green" if color == "red" else "red"

        cells[3].text = delta_str

        if color == "green":
            set_cell_shading(cells[3], GREEN_BG)
        elif color == "red":
            set_cell_shading(cells[3], RED_BG)

        # Alternate row shading
        if row_idx % 2 == 1:
            for cell in [cells[0], cells[1], cells[2]]:
                set_cell_shading(cell, LIGHT_GRAY)

    doc.add_paragraph()


def add_sp_outlook_section(doc: Document, company_info: CompanyInfo, ttm_revenue: float = None, unit: str = "dollars"):
    """Add the S&P/Moody's outlook section with editable placeholders."""
    doc.add_heading("Credit & Company Overview", level=2)

    table = doc.add_table(rows=8, cols=2)

    # Placeholder text for manual entry
    EDIT_PLACEHOLDER = "[EDIT]"

    rows_data = [
        ("S&P Rating", EDIT_PLACEHOLDER),
        ("S&P Outlook", EDIT_PLACEHOLDER),
        ("Moody's Rating", EDIT_PLACEHOLDER),
        ("Moody's Outlook", EDIT_PLACEHOLDER),
        ("HQ", f"{company_info.hq_city}, {company_info.hq_state}" if company_info.hq_city else EDIT_PLACEHOLDER),
        ("Public/Private?", "Public" if company_info.ticker else "Private"),
        ("Locations", EDIT_PLACEHOLDER),
        ("TTM Revenue", format_currency(ttm_revenue, unit=unit) if ttm_revenue else EDIT_PLACEHOLDER),
    ]

    for row_idx, (label, value) in enumerate(rows_data):
        cells = table.rows[row_idx].cells
        cells[0].text = label
        cells[1].text = value
        # Right-align the value
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Highlight placeholders in yellow for easy identification
        if value == EDIT_PLACEHOLDER:
            run = cells[1].paragraphs[0].runs[0]
            run.font.highlight_color = 7  # Yellow highlight

    doc.add_paragraph()


def add_company_overview(doc: Document, company_info: CompanyInfo, narrative: str, logo_path: Path = None):
    """Add company overview section with logo and narrative."""
    doc.add_heading("Company Overview", level=1)

    # Add logo if available
    if logo_path and logo_path.exists():
        doc.add_picture(str(logo_path), width=Inches(2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Company name and basic info
    doc.add_heading(company_info.name, level=2)

    info_para = doc.add_paragraph()
    info_para.add_run(f"Ticker: ").bold = True
    info_para.add_run(f"{company_info.ticker}  |  ")
    info_para.add_run(f"Sector: ").bold = True
    info_para.add_run(f"{company_info.sector}  |  ")
    info_para.add_run(f"Industry: ").bold = True
    info_para.add_run(company_info.industry)

    # Narrative with markdown formatting support
    if narrative:
        doc.add_paragraph()
        # Split narrative into lines and process each
        for line in narrative.split('\n'):
            if line.strip():
                add_formatted_paragraph(doc, line.strip())
            else:
                doc.add_paragraph()

    doc.add_paragraph()


def add_corporate_actions(doc: Document, actions: list[CorporateAction]):
    """Add key corporate actions section."""
    doc.add_heading("Key Corporate Actions", level=2)

    if not actions:
        doc.add_paragraph("No recent corporate actions available.")
        return

    for action in actions[:10]:  # Limit to 10 most recent
        para = doc.add_paragraph(style="List Bullet")
        para.add_run(f"{action.date}: ").bold = True
        para.add_run(f"{action.description}")
        if action.value:
            para.add_run(f" (${action.value:.2f})")

    doc.add_paragraph()


def add_ebitda_reconciliation(doc: Document, metrics: FinancialMetrics, fiscal_year: str, prior_year: str, unit: str = "dollars"):
    """Add EBITDA reconciliation table."""
    doc.add_heading("EBITDA Reconciliation", level=2)

    # Calculate implied D&A
    da_current = metrics.ebitda - metrics.operating_income if metrics.ebitda and metrics.operating_income else 0
    da_prior = metrics.ebitda_prior - metrics.operating_income_prior if metrics.ebitda_prior and metrics.operating_income_prior else 0

    # Calculate adjustments for Adjusted EBITDA
    adj_current = metrics.adjusted_ebitda - metrics.ebitda if metrics.adjusted_ebitda and metrics.ebitda else 0
    adj_prior = metrics.adjusted_ebitda_prior - metrics.ebitda_prior if metrics.adjusted_ebitda_prior and metrics.ebitda_prior else 0

    rows_data = [
        ("Net Income", metrics.net_income, metrics.net_income_prior),
        ("+ Interest Expense", 0, 0),  # Would need this from SEC data
        ("+ Income Tax Expense", 0, 0),  # Would need this from SEC data
        ("+ Depreciation & Amortization", da_current, da_prior),
        ("= EBITDA", metrics.ebitda, metrics.ebitda_prior),
        ("+ Adjustments", adj_current, adj_prior),
        ("= Adjusted EBITDA", metrics.adjusted_ebitda, metrics.adjusted_ebitda_prior),
    ]

    table = doc.add_table(rows=1 + len(rows_data), cols=3)
    table.style = "Table Grid"

    # Header
    header_cells = table.rows[0].cells
    headers = ["", fiscal_year[:4], prior_year[:4]]
    for i, header in enumerate(headers):
        header_cells[i].text = header
        set_cell_shading(header_cells[i], HEADER_BLUE)
        run = header_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

    for row_idx, (label, current, prior) in enumerate(rows_data):
        row = table.rows[row_idx + 1]
        cells = row.cells
        cells[0].text = label
        cells[1].text = format_currency(current, unit=unit)
        cells[2].text = format_currency(prior, unit=unit)

        # Bold the totals
        if label.startswith("="):
            for cell in cells:
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True

    doc.add_paragraph()


def add_detailed_statement_table(
    doc: Document,
    title: str,
    items: list,
    session: ExtractionSession,
    fiscal_year: str,
    prior_year: str,
    unit: str = "dollars",
):
    """Add a detailed financial statement table (Income Statement, Balance Sheet, or Cash Flow)."""
    doc.add_heading(title, level=2)

    # Filter to items with data
    visible = [item for item in items if item.metric_key in session.raw_values]
    if not visible:
        doc.add_paragraph(f"No {title.lower()} data available.")
        return

    table = doc.add_table(rows=1 + len(visible), cols=4)
    table.style = "Table Grid"
    table.autofit = False

    # Set column widths: wider label column, equal value columns
    col_widths = [Inches(2.8), Inches(1.4), Inches(1.4), Inches(1.4)]
    for col_idx, width in enumerate(col_widths):
        table.columns[col_idx].width = width

    # Header
    headers = [title, fiscal_year[:4], prior_year[:4], "Delta"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = col_widths[i]
        cell.text = header
        set_cell_shading(cell, HEADER_BLUE)
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)

    row_idx = 0
    for item in visible:
        row_idx += 1
        ev = session.raw_values[item.metric_key]
        row = table.rows[row_idx]
        cells = row.cells

        # Set widths on every row to prevent Word from auto-resizing
        for ci in range(4):
            cells[ci].width = col_widths[ci]

        # Label with indent
        label = ("    " if item.indent_level > 0 else "") + item.display_name
        cells[0].text = label
        for run in cells[0].paragraphs[0].runs:
            run.font.size = Pt(9)
            if item.is_bold:
                run.font.bold = True

        # Values
        for ci, val_text in enumerate([
            format_currency(ev.value, unit=unit),
            format_currency(ev.value_prior, unit=unit),
        ], start=1):
            cells[ci].text = val_text
            for run in cells[ci].paragraphs[0].runs:
                run.font.size = Pt(9)

        # Delta
        delta_str, color = format_delta(ev.value, ev.value_prior, unit=unit)
        cells[3].text = delta_str
        for run in cells[3].paragraphs[0].runs:
            run.font.size = Pt(9)
        if color == "green":
            set_cell_shading(cells[3], GREEN_BG)
        elif color == "red":
            set_cell_shading(cells[3], RED_BG)

        # Bold subtotal rows
        if item.is_bold:
            for cell in cells:
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True

        # Alternate row shading
        if row_idx % 2 == 0:
            for cell in [cells[0], cells[1], cells[2]]:
                set_cell_shading(cell, LIGHT_GRAY)

    doc.add_paragraph()


def add_verification_summary(doc: Document, verification: VerificationResult):
    """Add a verification summary section to the report."""
    doc.add_heading("Data Verification", level=2)

    # Summary paragraph
    active = [c for c in verification.checks if not c.skipped]
    passed = [c for c in active if c.passed]
    failed = [c for c in active if not c.passed]

    summary = f"{len(passed)} of {len(active)} checks passed"
    if verification.skip_count > 0:
        summary += f" ({verification.skip_count} skipped due to missing data)"

    para = doc.add_paragraph()
    run = para.add_run(summary)
    if len(failed) == 0:
        run.font.color.rgb = RGBColor(0, 128, 0)  # Green
    else:
        run.font.color.rgb = RGBColor(200, 0, 0)  # Red

    if failed:
        doc.add_paragraph()
        for check in failed:
            para = doc.add_paragraph(style="List Bullet")
            severity = "ERROR" if check.severity == "error" else "WARNING"
            para.add_run(f"[{severity}] ").bold = True
            para.add_run(f"{check.description} ({check.year}): {check.formula}")
            para.add_run(f" — Diff: {check.difference:,.0f}")

    doc.add_paragraph()


def generate_word_report(
    output_path,
    company_info: CompanyInfo,
    metrics: FinancialMetrics,
    ratios: FinancialRatios,
    fiscal_year_end: str,
    fiscal_year_end_prior: str,
    unit: str = "dollars",
    narrative: str = "",
    corporate_actions: list[CorporateAction] = None,
    logo_path: Path = None,
    sp_rating: str = "[EDIT]",
    sp_outlook: str = "[EDIT]",
    moodys_rating: str = "[EDIT]",
    moodys_outlook: str = "[EDIT]",
    session: ExtractionSession = None,
    verification: VerificationResult = None,
):
    """
    Generate a Word document financial report.

    Args:
        output_path: Path to save the .docx file, or BytesIO buffer
        company_info: Company information from Yahoo Finance
        metrics: Financial metrics
        ratios: Financial ratios
        fiscal_year_end: Current fiscal year end date (e.g., "2024-12-31")
        fiscal_year_end_prior: Prior fiscal year end date
        narrative: LLM-generated company narrative
        corporate_actions: List of corporate actions
        logo_path: Path to company logo image
        sp_rating: S&P rating (or [EDIT] placeholder)
        sp_outlook: S&P outlook (or [EDIT] placeholder)
        moodys_rating: Moody's rating (or [EDIT] placeholder)
        moodys_outlook: Moody's outlook (or [EDIT] placeholder)

    Returns:
        Path to the generated document (if output_path is a Path)

    Note:
        S&P/Moody's ratings, locations, and other manual data are marked with
        [EDIT] placeholders highlighted in yellow for the user to fill in.
    """
    doc = Document()

    # Title
    company_name = company_info.name if company_info else "Company"
    title = doc.add_heading(f"Financial Report: {company_name}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Company Overview
    if company_info:
        add_company_overview(doc, company_info, narrative, logo_path)

    # Corporate Actions
    add_corporate_actions(doc, corporate_actions or [])

    # S&P/Moody's section (with editable placeholders)
    if company_info:
        add_sp_outlook_section(doc, company_info, ttm_revenue=metrics.top_line_revenue, unit=unit)

    # Financial Statements Overview
    add_financial_overview_table(doc, metrics, fiscal_year_end, fiscal_year_end_prior, unit=unit)

    # Ratios
    add_ratios_table(doc, ratios, fiscal_year_end, fiscal_year_end_prior, unit=unit)

    # EBITDA Reconciliation
    add_ebitda_reconciliation(doc, metrics, fiscal_year_end, fiscal_year_end_prior, unit=unit)

    # Detailed Financial Statements (if session data available)
    if session:
        doc.add_page_break()
        doc.add_heading("Detailed Financial Statements", level=1)

        add_detailed_statement_table(
            doc, "Income Statement", INCOME_STATEMENT_ITEMS, session,
            fiscal_year_end, fiscal_year_end_prior, unit=unit,
        )
        add_detailed_statement_table(
            doc, "Balance Sheet", BALANCE_SHEET_ITEMS, session,
            fiscal_year_end, fiscal_year_end_prior, unit=unit,
        )
        add_detailed_statement_table(
            doc, "Cash Flow Statement", CASH_FLOW_ITEMS, session,
            fiscal_year_end, fiscal_year_end_prior, unit=unit,
        )

    # Verification Summary
    if verification:
        add_verification_summary(doc, verification)

    # Save - support both Path and BytesIO
    import io
    if isinstance(output_path, io.BytesIO):
        doc.save(output_path)
        return None
    else:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        return output_path
