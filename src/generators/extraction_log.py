"""Extraction log document generation."""

from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.calculators.metrics import FinancialMetrics
from src.calculators.ratios import FinancialRatios


def format_currency(value: float) -> str:
    """Format a currency value."""
    if value == 0:
        return "-"
    abs_val = abs(value)
    if abs_val >= 1e9:
        formatted = f"${abs_val / 1e9:,.2f} B"
    elif abs_val >= 1e6:
        formatted = f"${abs_val / 1e6:,.2f} M"
    else:
        formatted = f"${abs_val:,.0f}"
    return f"-{formatted}" if value < 0 else formatted


def format_percentage(value: float) -> str:
    """Format a decimal as percentage."""
    if value == 0:
        return "-"
    return f"{value * 100:.2f}%"


def format_ratio(value: float) -> str:
    """Format a ratio value."""
    if value == 0:
        return "-"
    return f"{value:.2f}x"


def generate_extraction_log(
    output_path: Path,
    ticker: str,
    company_name: str,
    fiscal_year_end: str,
    fiscal_year_end_prior: str,
    metrics: FinancialMetrics,
    ratios: FinancialRatios,
    notes: list[str],
    warnings: list[str],
) -> Path:
    """
    Generate a Word document with extraction notes, warnings, and raw data.

    This serves as an audit trail for the financial data extraction process.
    """
    doc = Document()

    # Title
    title = doc.add_heading(f"Extraction Log: {company_name} ({ticker})", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Fiscal Year End: {fiscal_year_end}")
    doc.add_paragraph(f"Prior Year End: {fiscal_year_end_prior}")

    # Warnings Section
    doc.add_heading("Warnings", level=1)
    if warnings:
        for warning in warnings:
            para = doc.add_paragraph(style="List Bullet")
            run = para.add_run(warning)
            run.font.color.rgb = RGBColor(192, 0, 0)  # Dark red
    else:
        doc.add_paragraph("No warnings.")

    # Extraction Notes Section
    doc.add_heading("Extraction Notes", level=1)
    if notes:
        for note in notes:
            doc.add_paragraph(note, style="List Bullet")
    else:
        doc.add_paragraph("No extraction notes.")

    # Raw Extracted Values - Metrics
    doc.add_heading("Extracted Financial Metrics", level=1)

    metrics_table = doc.add_table(rows=1, cols=4)
    metrics_table.style = "Table Grid"

    # Header
    header_cells = metrics_table.rows[0].cells
    header_cells[0].text = "Metric"
    header_cells[1].text = "Current Value"
    header_cells[2].text = "Prior Value"
    header_cells[3].text = "Delta"
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True

    # Metrics data
    metrics_data = [
        ("Tangible Net Worth", metrics.tangible_net_worth, metrics.tangible_net_worth_prior, True),
        ("Cash Balance", metrics.cash_balance, metrics.cash_balance_prior, True),
        ("Top Line Revenue", metrics.top_line_revenue, metrics.top_line_revenue_prior, True),
        ("Gross Profit", metrics.gross_profit, metrics.gross_profit_prior, True),
        ("Gross Profit Margin", metrics.gross_profit_margin, metrics.gross_profit_margin_prior, False),
        ("Operating Income", metrics.operating_income, metrics.operating_income_prior, True),
        ("Operating Income Margin", metrics.operating_income_margin, metrics.operating_income_margin_prior, False),
        ("EBITDA", metrics.ebitda, metrics.ebitda_prior, True),
        ("EBITDA Margin", metrics.ebitda_margin, metrics.ebitda_margin_prior, False),
        ("Adjusted EBITDA", metrics.adjusted_ebitda, metrics.adjusted_ebitda_prior, True),
        ("Adj. EBITDA Margin", metrics.adjusted_ebitda_margin, metrics.adjusted_ebitda_margin_prior, False),
        ("Net Income", metrics.net_income, metrics.net_income_prior, True),
        ("Net Income Margin", metrics.net_income_margin, metrics.net_income_margin_prior, False),
    ]

    for label, current, prior, is_currency in metrics_data:
        row = metrics_table.add_row()
        cells = row.cells
        cells[0].text = label

        if is_currency:
            cells[1].text = format_currency(current)
            cells[2].text = format_currency(prior)
            delta = current - prior
            cells[3].text = format_currency(delta) if delta != 0 else "-"
        else:
            cells[1].text = format_percentage(current)
            cells[2].text = format_percentage(prior)
            delta = current - prior
            cells[3].text = f"{delta * 100:+.2f}%" if delta != 0 else "-"

    doc.add_paragraph()

    # Raw Extracted Values - Ratios
    doc.add_heading("Extracted Financial Ratios", level=1)

    ratios_table = doc.add_table(rows=1, cols=4)
    ratios_table.style = "Table Grid"

    # Header
    header_cells = ratios_table.rows[0].cells
    header_cells[0].text = "Ratio"
    header_cells[1].text = "Current Value"
    header_cells[2].text = "Prior Value"
    header_cells[3].text = "Delta"
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True

    ratios_data = [
        ("Current Ratio", ratios.current_ratio, ratios.current_ratio_prior, False),
        ("Cash Ratio", ratios.cash_ratio, ratios.cash_ratio_prior, False),
        ("Debt-to-Equity", ratios.debt_to_equity, ratios.debt_to_equity_prior, False),
        ("EBITDA Interest Coverage", ratios.ebitda_interest_coverage, ratios.ebitda_interest_coverage_prior, False),
        ("Net Debt / EBITDA", ratios.net_debt_to_ebitda, ratios.net_debt_to_ebitda_prior, False),
        ("Net Debt / Adj. EBITDA", ratios.net_debt_to_adj_ebitda, ratios.net_debt_to_adj_ebitda_prior, False),
        ("Days Sales Outstanding", ratios.days_sales_outstanding, ratios.days_sales_outstanding_prior, True),
        ("Working Capital", ratios.working_capital, ratios.working_capital_prior, True),
        ("Return on Assets", ratios.return_on_assets, ratios.return_on_assets_prior, False),
        ("Return on Equity", ratios.return_on_equity, ratios.return_on_equity_prior, False),
    ]

    for label, current, prior, is_days_or_currency in ratios_data:
        row = ratios_table.add_row()
        cells = row.cells
        cells[0].text = label

        if "Working Capital" in label:
            cells[1].text = format_currency(current)
            cells[2].text = format_currency(prior)
            delta = current - prior
            cells[3].text = format_currency(delta) if delta != 0 else "-"
        elif is_days_or_currency:
            cells[1].text = f"{current:.1f}" if current else "-"
            cells[2].text = f"{prior:.1f}" if prior else "-"
            delta = current - prior
            cells[3].text = f"{delta:+.1f}" if delta != 0 else "-"
        else:
            cells[1].text = format_ratio(current)
            cells[2].text = format_ratio(prior)
            delta = current - prior
            cells[3].text = f"{delta:+.2f}x" if delta != 0 else "-"

    # Save
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    return output_path
